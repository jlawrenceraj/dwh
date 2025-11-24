from docx import Document
from docx.shared import Pt
import os

TECHNICAL_DESIGN = """
Technical Design: ETL CSV → PostgreSQL

Overview:
This document describes the technical design and architecture for the ETL tool in this workspace (CSV validation → PostgreSQL load, with error extraction + email).

Architecture Diagram (ASCII):

ETL Runner (CLI / Executable)
├─ `etl.py` / `src/main.py`  (ETLProcess)
│
├─ reads configuration
│   └─ `src/config/schema_config.yml`  (schema, table, DB, email, etc.)
│
├─ reads CSV
│   └─ `data/enterprise-survey-2024.csv`
│
├─ Validator
│   └─ `src/validator/csv_validator.py`  (validate_row, validate_csv)
│
├─ If errors: write `<csv>_errors.csv` → Emailer
│   ├─ `src/utils/email_utils.py`  (send_email attachments)
│   └─ SMTP server
│
└─ Loader
    ├─ `src/loader/postgres_loader.py`  (load_to_postgres(rows, db_config, table))
    └─ `src/utils/db_utils.py`  (get_connection, insert_rows using psycopg2)
        └─ PostgreSQL DB (schema: `etl.enterprise_survey` created via `src/create_table.sql`)

Sequence (run):
- CLI/Executable starts → loads YAML config (`load_config`)
- CSV path resolution (CLI arg or default)
- ETLProcess.validate() calls `validate_csv(schema, unique_fields)`:
  - per-row: required, length, type checks
  - global: duplicate detection on `unique_fields`
- Collects `valid_rows` and `error_rows`
- If `error_rows`:
  - writes `<csv>_errors.csv`
  - optional: sends email (SMTP info from YAML)
  - removes invalid rows from original CSV
- If `valid_rows`: calls loader to insert using `psycopg2` execute_values
- Exit with summary (counts, file paths, email status)

Key Files & Roles:
- `src/config/schema_config.yml` — single source of truth for schema, unique fields, table name, db_config, email
- `src/main.py` — `ETLProcess` class: validate, write_errors (and email), load
- `src/validator/csv_validator.py` — row + CSV validation logic
- `src/loader/postgres_loader.py` — parameterized table loader
- `src/utils/db_utils.py` — DB connection and bulk insert (execute_values)
- `src/utils/email_utils.py` — email helper (attachments)
- `src/create_table.sql` — DDL for the target table
- `pyinstaller.spec` + `etl.py` — packaging entry point

Configuration (YAML) — important keys (example):
- `schema:` per-column rules: `type`, `required`, `max_length`
- `unique_fields:` columns to dedupe
- `table_name:` `etl.enterprise_survey`
- `db_config:` host/port/user/password/dbname
- `email:` enabled/from/to/smtp

Operational Considerations:
- Idempotency: consider upsert or staging+merge for re-runs
- Transactions: single commit used; consider batching
- Error handling: write `<csv>_errors.csv` and email best-effort
- Secrets: do not store credentials in YAML for production; use environment variables or secret manager

Scalability & Performance:
- Bulk insert uses `execute_values`. For larger data consider `COPY FROM` or chunking.
- For very large CSVs, stream processing is recommended instead of loading into memory when writing errors.

Extensibility:
- Add connectors for S3 / other DBs
- Support schema evolution and default values
- Add unit tests and CI

"""


def build_docx(output_path=None):
    if output_path is None:
        output_path = os.path.join('src', 'docs', 'technical_design.docx')
    os.makedirs(os.path.dirname(output_path), exist_ok=True)

    doc = Document()
    doc.styles['Normal'].font.name = 'Calibri'
    doc.styles['Normal'].font.size = Pt(11)

    # Title
    doc.add_heading('Technical Design: ETL CSV to PostgreSQL', level=1)

    for para in TECHNICAL_DESIGN.strip().split('\n\n'):
        if para.startswith('Architecture Diagram'):
            doc.add_heading('Architecture Diagram', level=2)
            continue
        # For the ASCII diagram and code-like blocks, use monospaced paragraph
        if para.startswith('ETL Runner') or para.startswith('Sequence') or para.startswith('Key Files') or para.startswith('Configuration') or para.startswith('Operational') or para.startswith('Scalability') or para.startswith('Extensibility'):
            doc.add_paragraph(para)
        else:
            # preserve newlines inside the ASCII diagram by using runs
            if '\n' in para:
                p = doc.add_paragraph()
                for line in para.split('\n'):
                    run = p.add_run(line)
                    run.font.name = 'Consolas'
                    run.add_break()
            else:
                doc.add_paragraph(para)

    doc.save(output_path)
    print(f'Wrote: {output_path}')


if __name__ == '__main__':
    build_docx()
